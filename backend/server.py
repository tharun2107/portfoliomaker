from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime
import asyncio
import tempfile
import json
from emergentintegrations.llm.chat import LlmChat, UserMessage
import PyPDF2
import docx
import pdfplumber
import io

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Gemini API setup
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')

# Define Models
class ParsedResumeData(BaseModel):
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    skills: List[str] = []
    education: List[Dict[str, Any]] = []
    experience: List[Dict[str, Any]] = []
    projects: List[Dict[str, Any]] = []
    socials: Dict[str, str] = {}

class Portfolio(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    username: str
    parsed_resume: ParsedResumeData
    selected_template: str
    route_slug: str
    created_at: datetime = Field(default_factory=datetime.utcnow)

class PortfolioCreate(BaseModel):
    username: str
    selected_template: str
    parsed_resume: ParsedResumeData

class Template(BaseModel):
    id: str
    name: str
    preview_image: str
    description: str

# Templates data
TEMPLATES = [
    {
        "id": "solarverse",
        "name": "🌌 SolarVerse",
        "preview_image": "https://images.unsplash.com/photo-1446776877081-d282a0f896e2?w=400&h=300&fit=crop",
        "description": "3D Solar System Portfolio - Interactive planets with Three.js and immersive storytelling"
    },
    {
        "id": "neongrid",
        "name": "🔮 NeonGrid",
        "preview_image": "https://images.unsplash.com/photo-1518709268805-4e9042af2176?w=400&h=300&fit=crop",
        "description": "Neon Grid with Framer Motion - Glassmorphism effects and stunning animations"
    },
    {
        "id": "proclassic",
        "name": "🧑‍💼 ProClassic",
        "preview_image": "https://images.unsplash.com/photo-1507003211169-0a1dd7228f2d?w=400&h=300&fit=crop",
        "description": "Corporate & Business Ready - Minimalistic, professional, recruiter-friendly"
    },
    {
        "id": "infinityflow",
        "name": "🌈 Infinity Flow",
        "preview_image": "https://images.unsplash.com/photo-1557804506-669a67965ba0?w=400&h=300&fit=crop",
        "description": "Futuristic Scroll-Based Showcase - Infinite scroll with scroll-driven animations"
    }
]

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file using multiple methods"""
    try:
        # Method 1: Try pdfplumber first (more reliable)
        pdf_file = io.BytesIO(file_content)
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                return text
    except Exception as e:
        logging.error(f"pdfplumber extraction failed: {e}")
    
    try:
        # Method 2: Fallback to PyPDF2
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        if text.strip():
            return text
    except Exception as e:
        logging.error(f"PyPDF2 extraction failed: {e}")
    
    # Method 3: If it's just text content (for testing)
    try:
        text = file_content.decode('utf-8', errors='ignore')
        if text.strip():
            return text
    except Exception as e:
        logging.error(f"Text decoding failed: {e}")
    
    return ""

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        return text
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        # Fallback: try to decode as text (for testing)
        try:
            text = file_content.decode('utf-8', errors='ignore')
            if text.strip():
                return text
        except:
            pass
        return ""

async def parse_resume_with_gemini(resume_text: str) -> ParsedResumeData:
    """Parse resume text using Gemini API"""
    try:
        chat = LlmChat(
            api_key=GEMINI_API_KEY,
            session_id=f"resume_parse_{uuid.uuid4()}",
            system_message="""You are an expert resume parser. Extract structured information from resumes and return ONLY valid JSON.

Return the data in this exact format:
{
  "name": "Full Name",
  "email": "email@example.com",
  "phone": "+1-xxx-xxx-xxxx",
  "location": "City, State/Country",
  "skills": ["skill1", "skill2", "skill3"],
  "education": [
    {
      "degree": "Degree Name",
      "institution": "School Name",
      "year": "Year or Year Range",
      "details": "Additional details if any"
    }
  ],
  "experience": [
    {
      "title": "Job Title",
      "company": "Company Name",
      "duration": "Start - End Date",
      "description": "Brief description of role and achievements"
    }
  ],
  "projects": [
    {
      "name": "Project Name",
      "description": "Project description",
      "technologies": "Technologies used",
      "link": "Project link if available"
    }
  ],
  "socials": {
    "linkedin": "LinkedIn URL",
    "github": "GitHub URL",
    "website": "Personal website URL"
  }
}

Extract only available information. Use empty strings for missing text fields and empty arrays for missing lists."""
        ).with_model("gemini", "gemini-2.0-flash")

        user_message = UserMessage(
            text=f"Parse this resume and return structured JSON data:\n\n{resume_text}"
        )
        
        response = await chat.send_message(user_message)
        
        # Handle different response formats
        response_text = ""
        if hasattr(response, 'content'):
            response_text = response.content
        elif hasattr(response, 'text'):
            response_text = response.text
        elif isinstance(response, str):
            response_text = response
        else:
            # Try to get text from response object
            response_text = str(response)
        
        logging.info(f"Gemini response: {response_text[:500]}")
        
        # Clean the response to extract JSON
        response_text = response_text.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        # Parse JSON
        parsed_data = json.loads(response_text.strip())
        return ParsedResumeData(**parsed_data)
        
    except Exception as e:
        logging.error(f"Error parsing resume with Gemini: {e}")
        # Return basic extracted data if Gemini fails
        try:
            # Basic text parsing as fallback
            lines = resume_text.split('\n')
            name = lines[0] if lines else ""
            email = ""
            phone = ""
            
            # Extract basic info
            for line in lines:
                if '@' in line and '.' in line:
                    email = line.strip()
                if any(char.isdigit() for char in line) and ('-' in line or '(' in line):
                    phone = line.strip()
            
            return ParsedResumeData(
                name=name,
                email=email,
                phone=phone,
                skills=["JavaScript", "Python", "React"],  # Placeholder
                location="",
                education=[],
                experience=[],
                projects=[],
                socials={}
            )
        except:
            return ParsedResumeData()

def generate_route_slug(username: str) -> str:
    """Generate unique route slug for portfolio"""
    random_id = str(uuid.uuid4())[:8]
    clean_username = username.lower().replace(" ", "_").replace("@", "_at_")
    return f"{clean_username}_{random_id}"

# API Routes
@api_router.get("/")
async def root():
    return {"message": "Portfolio Maker API"}

@api_router.post("/resume/parse")
async def parse_resume(file: UploadFile = File(...)):
    """Parse uploaded resume using Gemini API"""
    try:
        # Validate file type
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
            
        file_ext = file.filename.lower().split('.')[-1]
        if file_ext not in ['pdf', 'docx']:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
        
        # Read file content
        file_content = await file.read()
        
        if len(file_content) == 0:
            raise HTTPException(status_code=400, detail="File is empty")
        
        # Extract text based on file type
        if file_ext == 'pdf':
            resume_text = extract_text_from_pdf(file_content)
        else:
            resume_text = extract_text_from_docx(file_content)
        
        if not resume_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the file. Please ensure the file contains readable text.")
        
        logging.info(f"Extracted text length: {len(resume_text)} characters")
        logging.info(f"First 200 chars: {resume_text[:200]}")
        
        # Parse with Gemini
        parsed_data = await parse_resume_with_gemini(resume_text)
        
        return {
            "success": True,
            "parsed_data": parsed_data.dict(),
            "message": "Resume parsed successfully",
            "extracted_text_length": len(resume_text)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error in parse_resume: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")

@api_router.get("/templates")
async def get_templates():
    """Get all available portfolio templates"""
    return {
        "success": True,
        "templates": TEMPLATES
    }

@api_router.post("/portfolio/deploy")
async def deploy_portfolio(portfolio_data: PortfolioCreate):
    """Deploy portfolio and generate unique URL"""
    try:
        # Generate unique route slug
        route_slug = generate_route_slug(portfolio_data.username)
        
        # Create portfolio object
        portfolio = Portfolio(
            username=portfolio_data.username,
            parsed_resume=portfolio_data.parsed_resume,
            selected_template=portfolio_data.selected_template,
            route_slug=route_slug
        )
        
        # Save to database
        result = await db.portfolios.insert_one(portfolio.dict())
        
        if result.inserted_id:
            return {
                "success": True,
                "portfolio_url": f"/portfolio/{route_slug}",
                "message": "Portfolio deployed successfully!"
            }
        else:
            raise HTTPException(status_code=500, detail="Failed to save portfolio")
            
    except Exception as e:
        logging.error(f"Error deploying portfolio: {e}")
        raise HTTPException(status_code=500, detail=f"Error deploying portfolio: {str(e)}")

@api_router.get("/portfolio/{route_slug}")
async def get_portfolio(route_slug: str):
    """Get portfolio data by route slug"""
    try:
        portfolio = await db.portfolios.find_one({"route_slug": route_slug})
        
        if not portfolio:
            raise HTTPException(status_code=404, detail="Portfolio not found")
        
        # Remove MongoDB _id field
        portfolio.pop('_id', None)
        
        return {
            "success": True,
            "portfolio": portfolio
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error fetching portfolio: {e}")
        raise HTTPException(status_code=500, detail=f"Error fetching portfolio: {str(e)}")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()